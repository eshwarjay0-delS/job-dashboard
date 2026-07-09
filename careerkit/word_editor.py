"""
word_editor.py — reads and edits .docx resumes in-place using python-docx.
Extracts plain text for Claude, then writes tailored content back.
"""

import copy
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


# ── Read ──────────────────────────────────────────────────────────────────────

def extract_text(filepath: str) -> str:
    """Return the full plain text of a .docx file."""
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


# ── Write (tailor in-place) ───────────────────────────────────────────────────

def apply_tailoring(src_path: str, dst_path: str, changes: dict) -> str:
    """
    Copy src → dst then apply `changes` dict:
      changes = {
        "summary": "new summary text",
        "skills":  ["skill1", "skill2", ...],
        "bullets": {"Old bullet text": "New bullet text", ...}
      }
    Returns dst_path on success.
    """
    doc = Document(src_path)

    summary_done = False
    skills_done = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # --- Summary (first long paragraph that isn't a heading) ---
        if (
            not summary_done
            and "summary" in changes
            and len(text) > 60
            and para.style.name not in ("Heading 1", "Heading 2", "Heading 3")
        ):
            _replace_para_text(para, changes["summary"])
            summary_done = True
            continue

        # --- Skills line ---
        if (
            not skills_done
            and "skills" in changes
            and re.search(r"\b(skills?|technologies|tech stack|tools)\b", text, re.I)
        ):
            skill_line = " • ".join(changes["skills"])
            _replace_para_text(para, skill_line)
            skills_done = True
            continue

        # --- Bullet replacements ---
        if "bullets" in changes:
            for old, new in changes["bullets"].items():
                if old.lower()[:40] in text.lower():
                    _replace_para_text(para, new)
                    break

    doc.save(dst_path)
    return dst_path


def _replace_para_text(para, new_text: str):
    """
    Replace all runs in a paragraph with new_text, preserving the font of
    the first run so formatting (bold, size, colour) is kept.
    """
    if not para.runs:
        para.text = new_text
        return

    # Grab first-run formatting
    first = para.runs[0]
    bold = first.bold
    font_name = first.font.name
    font_size = first.font.size

    # Clear all runs
    for run in para.runs:
        run.text = ""

    # Write into the first run
    first.text = new_text
    first.bold = bold
    if font_name:
        first.font.name = font_name
    if font_size:
        first.font.size = font_size


# ── PDF conversion ────────────────────────────────────────────────────────────

def convert_to_pdf(docx_path: str, pdf_path: str) -> str:
    """
    Convert a .docx to PDF.
    • Windows: uses Microsoft Word via COM automation.
    • Mac/Linux: falls back to LibreOffice.
    Returns pdf_path on success, raises on failure.
    """
    docx_path = str(Path(docx_path).resolve())
    pdf_path = str(Path(pdf_path).resolve())

    if sys.platform == "win32":
        _word_com_to_pdf(docx_path, pdf_path)
    else:
        _libreoffice_to_pdf(docx_path, pdf_path)

    return pdf_path


def _word_com_to_pdf(docx_path: str, pdf_path: str):
    import comtypes.client  # type: ignore

    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
    finally:
        word.Quit()


def _libreoffice_to_pdf(docx_path: str, pdf_path: str):
    out_dir = str(Path(pdf_path).parent)
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        check=True,
        capture_output=True,
    )
    # LibreOffice names the output after the source file
    generated = Path(out_dir) / (Path(docx_path).stem + ".pdf")
    if generated != Path(pdf_path):
        generated.rename(pdf_path)
